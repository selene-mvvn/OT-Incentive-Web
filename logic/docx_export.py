import io
import re
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

def strip_html_tags(text):
    clean = re.compile('<.*?>')
    text = re.sub(clean, '', text)
    text = text.replace('&nbsp;', ' ').replace('&ge;', '>=')
    return text

def apply_font_to_table(table, font_name="Times New Roman"):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name

def generate_project_executive_docx(df_project, project_name, period_label, analysis_text_vn, analysis_text_jp, total_hrs, total_cost, total_staff):
    doc = Document()
    
    # 1. Setup Page (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Header Title
    title = doc.add_heading('BÁO CÁO TỔNG QUAN DỰ ÁN', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.color.rgb = RGBColor(0, 176, 240)
    
    # Subtitle
    subtitle = doc.add_paragraph(f"Dự án: {project_name} | Thời gian: {period_label}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_heading('1. CHỈ SỐ HOẠT ĐỘNG CHÍNH (KPIs)', level=1)
    
    # KPI Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chỉ số'
    hdr_cells[1].text = 'Giá trị'
    
    row = table.add_row().cells
    row[0].text = 'Tổng thời gian OT'
    row[1].text = f"{total_hrs:,.1f} giờ"
    
    row = table.add_row().cells
    row[0].text = 'Tổng chi phí ước tính'
    row[1].text = f"{total_cost:,.0f} VNĐ"
    
    row = table.add_row().cells
    row[0].text = 'Tổng số nhân sự'
    row[1].text = f"{total_staff} người"
    
    apply_font_to_table(table)
    
    # AI Commentary
    doc.add_heading('2. PHÂN TÍCH TỪ AI', level=1)
    if analysis_text_vn:
        p = doc.add_paragraph(strip_html_tags(analysis_text_vn))
    if analysis_text_jp:
        p = doc.add_paragraph(strip_html_tags(analysis_text_jp))
        
    # Top Contributors
    doc.add_heading('3. NHÂN SỰ ĐÓNG GÓP CHÍNH (TOP 10)', level=1)
    if not df_project.empty:
        df_emp = df_project.groupby('employee_name')['ot_hours'].sum().reset_index()
        df_emp = df_emp.sort_values(by='ot_hours', ascending=False).head(10)
        
        emp_table = doc.add_table(rows=1, cols=3)
        emp_table.style = 'Light Grid Accent 1'
        hdr_cells = emp_table.rows[0].cells
        hdr_cells[0].text = 'STT'
        hdr_cells[1].text = 'Tên nhân sự'
        hdr_cells[2].text = 'Số giờ OT'
        
        for i, row in enumerate(df_emp.itertuples(), 1):
            r_cells = emp_table.add_row().cells
            r_cells[0].text = str(i)
            r_cells[1].text = str(row.employee_name)
            r_cells[2].text = f"{row.ot_hours:,.1f} giờ"
            
        apply_font_to_table(emp_table)
            
    # Add Charts using Matplotlib
    doc.add_page_break()
    doc.add_heading('4. PHÂN TÍCH BIỂU ĐỒ', level=1)
    
    if not df_project.empty:
        try:
            # 1. Employee Distribution Chart (fig_emp)
            staff_contrib = df_project.groupby('employee_name')['ot_hours'].sum().reset_index()
            staff_contrib.columns = ['employee_name', 'Hours']
            staff_contrib = staff_contrib.sort_values(by='Hours', ascending=True)
            
            plt.figure(figsize=(9, 4.5))
            bars = plt.barh(staff_contrib['employee_name'], staff_contrib['Hours'], color='#0284c7')
            plt.title('Biểu đồ Phân Bổ Số Giờ', fontname='Times New Roman', fontsize=14, fontweight='bold')
            plt.xlabel('Số giờ (h)', fontname='Times New Roman')
            
            # Add labels
            for bar in bars:
                plt.text(bar.get_width(), bar.get_y() + bar.get_height()/2, f'{bar.get_width():,.1f} h', 
                         va='center', ha='right', color='white', fontweight='bold', fontname='Times New Roman')
            
            plt.tight_layout()
            img_io1 = io.BytesIO()
            plt.savefig(img_io1, format='png', dpi=150)
            img_io1.seek(0)
            plt.close()
            
            p_img1 = doc.add_paragraph()
            doc.add_picture(img_io1, width=Inches(6.5))
            
            # 2. Time Trend Chart (fig_time)
            time_df = df_project.groupby('ot_date')['ot_hours'].sum().reset_index()
            time_df['_sort_dt'] = pd.to_datetime(time_df['ot_date'], format='%d/%m/%Y', errors='coerce')
            time_df = time_df.sort_values(by='_sort_dt', ascending=True)
            
            plt.figure(figsize=(9, 4.5))
            bars = plt.bar(time_df['ot_date'], time_df['ot_hours'], color='#ca8a04')
            plt.title('Biểu đồ Xu Hướng Bù Đắp OT', fontname='Times New Roman', fontsize=14, fontweight='bold')
            plt.xlabel('Ngày OT', fontname='Times New Roman')
            plt.ylabel('Số giờ (h)', fontname='Times New Roman')
            plt.xticks(rotation=45, ha='right')
            
            for bar in bars:
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height(), f'{bar.get_height():,.1f} h', 
                         va='bottom', ha='center', color='black', fontweight='bold', fontname='Times New Roman')
                         
            plt.tight_layout()
            img_io2 = io.BytesIO()
            plt.savefig(img_io2, format='png', dpi=150)
            img_io2.seek(0)
            plt.close()
            
            p_img2 = doc.add_paragraph()
            doc.add_picture(img_io2, width=Inches(6.5))
            
        except Exception as e:
            doc.add_paragraph(f"Không thể render biểu đồ: {e}")
            
    # Fix Times New Roman fallback in styles
    for style in doc.styles:
        if hasattr(style, 'font') and style.font.name is not None:
            style.font.name = 'Times New Roman'
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
